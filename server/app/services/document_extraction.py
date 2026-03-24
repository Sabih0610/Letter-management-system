from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re
from email import policy
from email.parser import BytesParser

from docx import Document
from pypdf import PdfReader

try:
    from PIL import Image
except ImportError:  # pragma: no cover
    Image = None  # type: ignore[assignment]

try:
    import pytesseract
except ImportError:  # pragma: no cover
    pytesseract = None  # type: ignore[assignment]

try:
    import pypdfium2
except ImportError:  # pragma: no cover
    pypdfium2 = None  # type: ignore[assignment]


@dataclass
class ExtractionResult:
    source_type: str
    status: str
    text: str | None = None
    error: str | None = None


def _normalize_text(value: str) -> str:
    value = value.replace("\u0000", " ")
    value = re.sub(r"\r\n?", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    value = re.sub(r"[ \t]{2,}", " ", value)
    return value.strip()


def _extract_pdf(content: bytes) -> str:
    reader = PdfReader(BytesIO(content))
    parts: list[str] = []
    for page in reader.pages:
        parts.append(page.extract_text() or "")
    return _normalize_text("\n\n".join(parts))


def _extract_docx(content: bytes) -> str:
    doc = Document(BytesIO(content))
    parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
    return _normalize_text("\n".join(parts))


def _extract_email(content: bytes) -> str:
    msg = BytesParser(policy=policy.default).parsebytes(content)
    body = ""
    preferred = msg.get_body(preferencelist=("plain", "html"))
    if preferred:
        body = preferred.get_content()
    elif msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                body = part.get_content()
                break
    else:
        body = msg.get_content()

    if not body:
        return ""
    body = re.sub(r"<[^>]+>", " ", body)
    return _normalize_text(body)


def _extract_text_plain(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return _normalize_text(content.decode(encoding))
        except UnicodeDecodeError:
            continue
    return _normalize_text(content.decode("utf-8", errors="ignore"))


def _ocr_image_bytes(content: bytes) -> str:
    if not Image or not pytesseract:
        return ""
    image = Image.open(BytesIO(content))
    text = pytesseract.image_to_string(image)
    return _normalize_text(text)


def _ocr_pdf_bytes(content: bytes, max_pages: int = 5) -> str:
    if not pypdfium2 or not pytesseract:
        return ""
    pdf = pypdfium2.PdfDocument(BytesIO(content))
    pages = min(len(pdf), max_pages)
    chunks: list[str] = []
    for idx in range(pages):
        page = pdf[idx]
        bitmap = page.render(scale=2.0)
        pil_image = bitmap.to_pil()
        chunks.append(pytesseract.image_to_string(pil_image))
    return _normalize_text("\n\n".join(chunks))


def extract_text_from_file(filename: str, mime_type: str | None, content: bytes) -> ExtractionResult:
    ext = Path(filename).suffix.lower()
    source_type = ext.lstrip(".") or "unknown"

    try:
        if ext == ".pdf":
            text = _extract_pdf(content)
            if len(text) < 80:
                ocr_text = _ocr_pdf_bytes(content)
                if len(ocr_text) > len(text):
                    text = ocr_text
            return ExtractionResult(source_type="pdf", status="completed", text=text)

        if ext in {".docx", ".doc"} or (mime_type and "word" in mime_type):
            return ExtractionResult(source_type="docx", status="completed", text=_extract_docx(content))

        if ext in {".eml", ".msg"} or (mime_type and "message" in mime_type):
            return ExtractionResult(source_type="email", status="completed", text=_extract_email(content))

        if ext in {".txt", ".md"} or (mime_type and mime_type.startswith("text/")):
            return ExtractionResult(source_type="text", status="completed", text=_extract_text_plain(content))

        if ext in {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp"} or (mime_type and mime_type.startswith("image/")):
            text = _ocr_image_bytes(content)
            return ExtractionResult(source_type="image_ocr", status="completed", text=text)

        return ExtractionResult(source_type=source_type, status="unsupported", text="")
    except Exception as exc:  # pragma: no cover
        return ExtractionResult(source_type=source_type, status="failed", text="", error=str(exc))

